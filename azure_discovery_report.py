#!/usr/bin/env python3
"""
Azure subscription discovery -> client meeting DOCX.

Requirements:
  - Azure CLI installed and authenticated (az login)
  - Python 3.10+
  - python-docx
  - matplotlib (for the architecture figure)

The script performs read-only Azure CLI discovery and generates a client-facing
DOCX report modeled on the project's current-state/discovery meeting document.
It intentionally reports incomplete discovery as "Not verified" instead of
assuming that an empty query means a resource does not exist.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
except Exception:
    plt = None


# -----------------------------
# Helpers
# -----------------------------

ACCENT = "1F4E79"
ACCENT_DARK = "173A5B"
ACCENT_LIGHT = "EAF2F8"
GRAY = "5B6573"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
DARK = "20252B"
WARNING = "FFF4CC"
SUCCESS = "EAF5EA"


def run_az(args: list[str], *, timeout: int = 120) -> tuple[Any, str | None]:
    """Run az command, parse JSON when possible, and return (data, error)."""
    cmd = ["az", *args]
    try:
        p = subprocess.run(cmd, text=True, capture_output=True, timeout=timeout)
    except FileNotFoundError:
        return None, "Azure CLI (az) was not found in PATH."
    except subprocess.TimeoutExpired:
        return None, f"Command timed out after {timeout}s: {' '.join(cmd)}"

    if p.returncode != 0:
        msg = (p.stderr or p.stdout or "Unknown Azure CLI error").strip()
        return None, msg

    out = p.stdout.strip()
    if not out:
        return [], None
    try:
        return json.loads(out), None
    except json.JSONDecodeError:
        return out, None


def safe_text(value: Any, default: str = "Not verified") -> str:
    if value is None or value == "":
        return default
    if isinstance(value, (dict, list)):
        return json.dumps(value, indent=2, default=str)
    return str(value)


def q(data: Any, path: str, default: Any = None) -> Any:
    """Tiny dotted-path getter for dict/list payloads."""
    cur = data
    for part in path.split("."):
        if isinstance(cur, dict):
            cur = cur.get(part, default)
        else:
            return default
    return cur


def first(items: list[Any], default: Any = None) -> Any:
    return items[0] if items else default


def xml_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold=False, color=DARK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        xml_shading(c, ACCENT)
        set_cell_text(c, h, bold=True, color=WHITE, size=9.2)
        if widths:
            c.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if widths:
                cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], val, size=9.2)
    # Repeat header row.
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, title: str, text: str, fill: str = ACCENT_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    c = table.cell(0, 0)
    xml_shading(c, fill)
    cell_margins(c, top=150, start=180, bottom=150, end=180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    p2 = c.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    rr = p2.add_run()
    rr.font.name = "Aptos"
    rr._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    rr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    title.paragraph_format.space_after = Pt(5)

    for sname in ("Heading 1", "Heading 2"):
        s = styles[sname]
        s.font.name = "Aptos Display"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].paragraph_format.space_before = Pt(11)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)

    if "Callout" not in styles:
        styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout = styles["Callout"]
    callout.font.name = "Aptos"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(DARK)

    # Footer.
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Azure Current State & Project Discovery")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def architecture_figure(path: Path, current_subnet: str, east_cidr: str, hub_cidr: str) -> None:
    if plt is None:
        return
    fig = plt.figure(figsize=(8.5, 4.8), dpi=180)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    def box(x, y, w, h, title, body, fill="#EAF2F8", edge="#1F4E79"):
        patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.6,rounding_size=2.0", linewidth=1.4,
                               edgecolor=edge, facecolor=fill)
        ax.add_patch(patch)
        ax.text(x + 2, y + h - 5, title, fontsize=10, fontweight="bold", color="#173A5B", va="top")
        ax.text(x + 2, y + h - 11, body, fontsize=8.3, color="#20252B", va="top", linespacing=1.25)

    box(3, 14, 27, 72, "Existing West US 3", f"VN-Accelevation\n{current_subnet}\n\nExisting AVD / AAD DS\nFSLogix / SQL / support", fill="#F3F5F7", edge="#6C757D")
    box(37, 39, 26, 40, "Azure Virtual WAN", f"Standard\nEast US Virtual Hub\n{hub_cidr}\n\nCisco C8000V x2", fill="#EAF5EA", edge="#3B6B48")
    box(71, 14, 26, 72, "New East US VNet", f"{east_cidr}\n\nISE\nIdentity / DNS\nPKI\nManagement\nWorkloads", fill="#EAF2F8", edge="#1F4E79")

    ax.add_patch(FancyArrowPatch((30, 50), (37, 58), arrowstyle="-|>", mutation_scale=14, linewidth=1.8, color="#6C757D"))
    ax.add_patch(FancyArrowPatch((63, 58), (71, 58), arrowstyle="-|>", mutation_scale=14, linewidth=1.8, color="#1F4E79"))

    ax.text(50, 5, "Working target architecture; existing West US 3 remains isolated from the new project network.",
            fontsize=8.3, color="#5B6573", ha="center")
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


@dataclass
class Discovery:
    account: dict[str, Any]
    groups: list[dict[str, Any]]
    resources: list[dict[str, Any]]
    vnets: list[dict[str, Any]]
    nsgs: list[dict[str, Any]]
    public_ips: list[dict[str, Any]]
    nic_rows: list[dict[str, Any]]
    private_endpoints: list[dict[str, Any]]
    private_dns_zones: list[dict[str, Any]]
    vms: list[dict[str, Any]]
    vm_details: dict[str, dict[str, Any]]
    vm_nics: dict[str, list[dict[str, Any]]]
    vm_disks: dict[str, list[dict[str, Any]]]
    vnet_detail: dict[str, Any] | None
    nsg_rules: dict[str, list[dict[str, Any]]]
    subscription_resource_checks: dict[str, Any]
    errors: list[str]



def discover_vm_details(vms: list[dict[str, Any]], errors: list[str]) -> tuple[dict[str, dict[str, Any]], dict[str, list[dict[str, Any]]], dict[str, list[dict[str, Any]]]]:
    """Collect per-VM detail and attached NIC/disk resources using read-only queries."""
    details: dict[str, dict[str, Any]] = {}
    vm_nics: dict[str, list[dict[str, Any]]] = {}
    vm_disks: dict[str, list[dict[str, Any]]] = {}

    for vm in vms:
        vm_id = vm.get("id")
        vm_name = vm.get("name")
        rg = vm.get("resourceGroup")
        if not vm_id or not vm_name or not rg:
            continue
        detail, err = run_az(["vm", "show", "--ids", vm_id, "-o", "json"])
        if err:
            errors.append(f"VM detail {rg}/{vm_name}: {err}")
            continue
        if not isinstance(detail, dict):
            continue
        details[vm_name] = detail

        nic_ids = q(detail, "networkProfile.networkInterfaces", []) or []
        nic_rows: list[dict[str, Any]] = []
        for nic_ref in nic_ids:
            nic_id = nic_ref.get("id") if isinstance(nic_ref, dict) else nic_ref
            if not nic_id:
                continue
            nic, nerr = run_az(["network", "nic", "show", "--ids", nic_id, "-o", "json"])
            if nerr:
                errors.append(f"VM NIC {vm_name}/{nic_id}: {nerr}")
                continue
            if isinstance(nic, dict):
                nic_rows.append(nic)
        vm_nics[vm_name] = nic_rows

        disk_rows: list[dict[str, Any]] = []
        storage = q(detail, "storageProfile", {}) or {}
        os_disk = storage.get("osDisk") if isinstance(storage, dict) else None
        if isinstance(os_disk, dict):
            managed = os_disk.get("managedDisk") or {}
            disk_id = managed.get("id")
            if disk_id:
                disk, derr = run_az(["disk", "show", "--ids", disk_id, "-o", "json"])
                if derr:
                    errors.append(f"VM OS disk {vm_name}/{disk_id}: {derr}")
                elif isinstance(disk, dict):
                    disk_rows.append(disk)
        for data_disk in (storage.get("dataDisks", []) if isinstance(storage, dict) else []):
            managed = (data_disk or {}).get("managedDisk") or {}
            disk_id = managed.get("id")
            if disk_id:
                disk, derr = run_az(["disk", "show", "--ids", disk_id, "-o", "json"])
                if derr:
                    errors.append(f"VM data disk {vm_name}/{disk_id}: {derr}")
                elif isinstance(disk, dict):
                    disk_rows.append(disk)
        vm_disks[vm_name] = disk_rows

    return details, vm_nics, vm_disks


def extract_vm_network_rows(
    vm: dict[str, Any],
    vm_detail: dict[str, Any] | None,
    nic_rows: list[dict[str, Any]],
    public_ip_map: dict[str, str] | None = None,
) -> tuple[str, str, str, str, str]:
    """Return private IPs, public IPs, VNet/subnet, NSGs, and NIC names for a VM."""
    public_ip_map = public_ip_map or {}
    private_ips: list[str] = []
    public_ips: list[str] = []
    subnets: list[str] = []
    nsgs: list[str] = []
    nic_names: list[str] = []

    for nic in nic_rows:
        nic_names.append(str(nic.get("name", "Unknown")))
        if nic.get("networkSecurityGroup", {}).get("id"):
            nsgs.append(str(nic["networkSecurityGroup"]["id"]).split("/")[-1])
        for ipcfg in nic.get("ipConfigurations", []) or []:
            if ipcfg.get("privateIPAddress"):
                private_ips.append(str(ipcfg["privateIPAddress"]))
            pip = (ipcfg.get("publicIPAddress") or {}).get("id")
            if pip:
                public_ips.append(public_ip_map.get(str(pip), str(pip).split("/")[-1]))
            subnet_id = (ipcfg.get("subnet") or {}).get("id")
            if subnet_id:
                parts = str(subnet_id).split("/")
                if "subnets" in parts:
                    idx = parts.index("subnets")
                    if idx + 1 < len(parts):
                        subnets.append(parts[idx + 1])
    vnet_name = "Not verified"
    if nic_rows and (nic_rows[0].get("ipConfigurations") or []):
        first_subnet = ((nic_rows[0].get("ipConfigurations") or [])[0].get("subnet") or {}).get("id")
        if first_subnet:
            parts = str(first_subnet).split("/")
            if "virtualNetworks" in parts:
                idx = parts.index("virtualNetworks")
                if idx + 1 < len(parts):
                    vnet_name = parts[idx + 1]

    return (
        ", ".join(dict.fromkeys(private_ips)) or "None",
        ", ".join(dict.fromkeys(public_ips)) or "None",
        f"{vnet_name} / {', '.join(dict.fromkeys(subnets)) or 'Not verified'}",
        ", ".join(dict.fromkeys(nsgs)) or "None/Not verified",
        ", ".join(dict.fromkeys(nic_names)) or "None",
    )

def discover(subscription_id: str | None = None) -> Discovery:
    errors: list[str] = []

    account_args = ["account", "show", "-o", "json"]
    if subscription_id:
        # Account selection is explicit and non-destructive.
        _, err = run_az(["account", "set", "--subscription", subscription_id])
        if err:
            errors.append(f"Could not set subscription: {err}")
    account, err = run_az(account_args)
    if err:
        errors.append(f"az account show: {err}")
        account = {}

    groups, err = run_az(["group", "list", "-o", "json"])
    if err:
        errors.append(f"az group list: {err}")
        groups = []

    resources, err = run_az(["resource", "list", "-o", "json"])
    if err:
        errors.append(f"az resource list: {err}")
        resources = []

    vnets, err = run_az(["network", "vnet", "list", "-o", "json"])
    if err:
        errors.append(f"az network vnet list: {err}")
        vnets = []

    nsgs, err = run_az(["network", "nsg", "list", "-o", "json"])
    if err:
        errors.append(f"az network nsg list: {err}")
        nsgs = []

    public_ips, err = run_az(["network", "public-ip", "list", "-o", "json"])
    if err:
        errors.append(f"az network public-ip list: {err}")
        public_ips = []

    nic_rows, err = run_az(["network", "nic", "list", "-o", "json"])
    if err:
        errors.append(f"az network nic list: {err}")
        nic_rows = []

    private_endpoints, err = run_az(["network", "private-endpoint", "list", "-o", "json"])
    if err:
        errors.append(f"az network private-endpoint list: {err}")
        private_endpoints = []

    private_dns_zones, err = run_az(["network", "private-dns", "zone", "list", "-o", "json"])
    if err:
        errors.append(f"az network private-dns zone list: {err}")
        private_dns_zones = []

    vnet_detail = None
    for vnet in vnets:
        name = vnet.get("name")
        rg = vnet.get("resourceGroup")
        if name and rg:
            detail, derr = run_az(["network", "vnet", "show", "-g", rg, "-n", name, "-o", "json"])
            if derr:
                errors.append(f"VNet detail {rg}/{name}: {derr}")
            elif isinstance(detail, dict):
                if vnet_detail is None:
                    vnet_detail = detail

    nsg_rules: dict[str, list[dict[str, Any]]] = {}
    for nsg in nsgs:
        rg = nsg.get("resourceGroup")
        name = nsg.get("name")
        if rg and name:
            rules, rerr = run_az(["network", "nsg", "rule", "list", "-g", rg, "--nsg-name", name, "-o", "json"])
            if rerr:
                errors.append(f"NSG rules {rg}/{name}: {rerr}")
            else:
                nsg_rules[name] = rules if isinstance(rules, list) else []

    # Virtual machine inventory and per-VM attached resource discovery.
    vms, err = run_az(["vm", "list", "-d", "-o", "json"])
    if err:
        errors.append(f"az vm list: {err}")
        vms = []
    vm_details, vm_nics, vm_disks = discover_vm_details(vms if isinstance(vms, list) else [], errors)

    # Subscription-wide resource type discovery, avoiding the older az network list limitation.
    type_map = {
        "routeTables": "Microsoft.Network/routeTables",
        "virtualNetworkGateways": "Microsoft.Network/virtualNetworkGateways",
        "azureFirewalls": "Microsoft.Network/azureFirewalls",
        "natGateways": "Microsoft.Network/natGateways",
        "applicationGateways": "Microsoft.Network/applicationGateways",
        "expressRouteCircuits": "Microsoft.Network/expressRouteCircuits",
        "connections": "Microsoft.Network/connections",
        "localNetworkGateways": "Microsoft.Network/localNetworkGateways",
        "loadBalancers": "Microsoft.Network/loadBalancers",
    }
    checks: dict[str, Any] = {}
    for key, rtype in type_map.items():
        data, rerr = run_az(["resource", "list", "--resource-type", rtype, "-o", "json"])
        if rerr:
            errors.append(f"Subscription discovery {rtype}: {rerr}")
            checks[key] = None
        else:
            checks[key] = data if isinstance(data, list) else []

    return Discovery(
        account=account or {},
        groups=groups if isinstance(groups, list) else [],
        resources=resources if isinstance(resources, list) else [],
        vnets=vnets if isinstance(vnets, list) else [],
        nsgs=nsgs if isinstance(nsgs, list) else [],
        public_ips=public_ips if isinstance(public_ips, list) else [],
        nic_rows=nic_rows if isinstance(nic_rows, list) else [],
        private_endpoints=private_endpoints if isinstance(private_endpoints, list) else [],
        private_dns_zones=private_dns_zones if isinstance(private_dns_zones, list) else [],
        vms=vms if isinstance(vms, list) else [],
        vm_details=vm_details,
        vm_nics=vm_nics,
        vm_disks=vm_disks,
        vnet_detail=vnet_detail,
        nsg_rules=nsg_rules,
        subscription_resource_checks=checks,
        errors=errors,
    )


def resources_by_type(d: Discovery, contains: str) -> list[dict[str, Any]]:
    target = contains.lower()
    if target == "microsoft.compute/virtualmachines":
        return [r for r in d.resources if str(r.get("type", "")).lower() == target]
    return [r for r in d.resources if target in str(r.get("type", "")).lower()]


def resource_counts_by_type(d: Discovery) -> list[list[str]]:
    counts: dict[str, int] = {}
    for r in d.resources:
        t = str(r.get("type", "Unknown"))
        counts[t] = counts.get(t, 0) + 1
    rows = []
    for t, c in sorted(counts.items(), key=lambda x: (-x[1], x[0])):
        rows.append([t, str(c)])
    return rows[:18]


def rg_names(d: Discovery) -> list[str]:
    return sorted(str(g.get("name")) for g in d.groups if g.get("name"))


def build_report(d: Discovery, output_path: Path, proposed_region: str = "East US") -> None:
    doc = Document()
    configure_document(doc)

    # Front matter.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("AZURE")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    p = doc.add_paragraph(style="Title")
    p.add_run("Current State & Project Discovery")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Client Working Session")
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GRAY)

    sub_id = q(d.account, "id", "Unknown")
    tenant_id = q(d.account, "tenantId", "Unknown")
    tenant_name = q(d.account, "name", "Unknown")
    region = proposed_region
    existing_vnets = d.vnets
    primary_vnet = first(existing_vnets, {}) or {}
    current_vnet_name = primary_vnet.get("name", "Not verified")
    current_vnet_location = primary_vnet.get("location", "Not verified")
    current_vnet_space = first(q(primary_vnet, "addressSpace.addressPrefixes", []), "Not verified")

    doc.add_heading("Executive takeaway", level=1)
    vnet_phrase = f"The current environment is concentrated in {current_vnet_location} and uses {current_vnet_space}."
    if current_vnet_name != "Not verified":
        vnet_phrase += f" The primary discovered VNet is {current_vnet_name}."
    doc.add_paragraph(
        f"The existing Azure subscription is active and in use. {vnet_phrase} "
        "The discovery shows existing identity, virtual desktop, Windows workload, storage/private endpoint, backup, and monitoring services. "
        f"The new project should be built in a separate {region} VNet and connected through the planned Azure Virtual WAN design rather than modifying the existing production VNet.",
        style="Callout"
    )
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph(f"Discovery date: {date.today().strftime('%B %-d, %Y')}")
    p.paragraph_format.space_after = Pt(2)

    doc.add_heading("1. What exists today", level=1)
    doc.add_paragraph(
        "The following reflects read-only Azure CLI discovery of the active subscription. "
        "The report intentionally distinguishes observed resources from items that could not be verified."
    )

    doc.add_heading("2. Current environment at a glance", level=1)
    at_glance_rows = [
        ["Tenant", safe_text(tenant_name)],
        ["Tenant ID", safe_text(tenant_id)],
        ["Subscription", safe_text(q(d.account, "name"))],
        ["Subscription ID", safe_text(sub_id)],
        ["Existing resource groups", ", ".join(rg_names(d)) or "Not verified"],
        ["New project region", region],
        ["Existing VNets discovered", str(len(d.vnets))],
        ["Existing public IPs", str(len(d.public_ips))],
        ["Existing VMs", str(len(d.vms))],
        ["Existing private endpoints", str(len(d.private_endpoints))],
    ]
    add_table(doc, ["Area", "Current state"], at_glance_rows, [2.0, 4.9])

    # Resource/service summary.
    doc.add_heading("3. Existing services discovered", level=1)
    service_rows: list[list[str]] = []
    patterns = [
        ("Virtual Network", "Microsoft.Network/virtualNetworks", "Existing Azure network boundary."),
        ("Virtual Machines", "Microsoft.Compute/virtualMachines", "Existing server/session-host footprint."),
        ("Azure Virtual Desktop", "Microsoft.DesktopVirtualization", "Existing AVD service components."),
        ("Azure AD Domain Services", "Microsoft.AAD/DomainServices", "Existing managed identity/domain services."),
        ("Storage", "Microsoft.Storage", "Existing storage / diagnostics / FSLogix-related resources."),
        ("Private Endpoints", "Microsoft.Network/privateEndpoints", "Existing private connectivity to PaaS services."),
        ("Private DNS", "Microsoft.Network/privateDnsZones", "Existing private name-resolution dependencies."),
        ("Recovery Services", "Microsoft.RecoveryServices", "Existing backup/recovery services."),
        ("Log Analytics", "Microsoft.OperationalInsights", "Existing monitoring workspace(s)."),
        ("SQL VM", "Microsoft.SqlVirtualMachine", "Existing SQL virtual machine integration."),
    ]
    for label, ptn, why in patterns:
        items = resources_by_type(d, ptn)
        if items:
            names = ", ".join(str(x.get("name")) for x in items[:8])
            if len(items) > 8:
                names += f" (+{len(items)-8} more)"
            service_rows.append([label, names, why])
    if not service_rows:
        service_rows.append(["Resource inventory", "No resources returned", "Discovery output was empty; verify permissions and subscription context."])
    add_table(doc, ["Service / resource", "Current observation", "Why it matters"], service_rows, [1.55, 3.0, 2.35])

    doc.add_paragraph("Resource type counts from the subscription:")
    add_table(doc, ["Resource type", "Count"], resource_counts_by_type(d), [5.8, 1.1])

    # VM inventory.
    doc.add_heading("4. Virtual machine inventory", level=1)
    doc.add_paragraph(
        "All virtual machines returned by Azure CLI are listed below. The inventory includes compute state, operating system, Azure size, "
        "private/public IP addressing, VNet/subnet placement, NSG associations, NICs, disks, and detected VM extensions."
    )
    vm_network_rows: list[list[str]] = []
    vm_resource_rows: list[list[str]] = []
    public_ip_map = {}
    for pip in d.public_ips:
        pid = pip.get("id")
        if pid and pip.get("ipAddress"):
            public_ip_map[str(pid)] = str(pip.get("ipAddress"))
    for vm in sorted(d.vms, key=lambda x: str(x.get("name", "")).lower()):
        name = str(vm.get("name", "Unknown"))
        detail = d.vm_details.get(name, {})
        nic_rows = d.vm_nics.get(name, [])
        disks = d.vm_disks.get(name, [])
        private_ip, public_ip, placement, nsgs, nics = extract_vm_network_rows(vm, detail, nic_rows, public_ip_map)
        os_name = safe_text(q(detail, "storageProfile.osDisk.osType", vm.get("storageProfile", {}).get("osDisk", {}).get("osType", "Not verified")))
        size = safe_text(q(detail, "hardwareProfile.vmSize", vm.get("hardwareProfile", {}).get("vmSize", "Not verified")))
        power = safe_text(vm.get("powerState"), "Not verified")
        rg = safe_text(vm.get("resourceGroup"))
        location = safe_text(vm.get("location"))
        zone = ", ".join(str(z) for z in (detail.get("zones") or [])) or "Not specified"
        vm_network_rows.append([name, f"{rg} / {location}", power, os_name, size, private_ip, public_ip, placement])

        disk_names: list[str] = []
        for disk in disks:
            disk_names.append(f"{disk.get('name', 'Unknown')} ({disk.get('diskSizeGB', '?')} GB, {disk.get('sku', {}).get('name', 'SKU ?')})")
        extension_names = []
        prefix = name.lower() + "/"
        for resource in d.resources:
            rname = str(resource.get("name", ""))
            rtype = str(resource.get("type", ""))
            if rname.lower().startswith(prefix) and rtype.lower().endswith("/extensions"):
                extension_names.append(rname.split("/", 1)[1] if "/" in rname else rname)
        vm_resource_rows.append([
            name,
            nics,
            nsgs,
            ", ".join(disk_names) or "None/Not verified",
            ", ".join(sorted(set(extension_names))) or "None/Not verified",
            zone,
        ])

    if vm_network_rows:
        add_table(doc, ["VM", "RG / Region", "State", "OS", "Size", "Private IP", "Public IP", "VNet / Subnet"], vm_network_rows,
                  [1.0, 0.95, 0.62, 0.5, 0.72, 0.72, 0.72, 1.77])
        doc.add_paragraph("Attached VM resources:")
        add_table(doc, ["VM", "NIC(s)", "NSG(s)", "Disk(s)", "Extensions", "Zone"], vm_resource_rows,
                  [0.85, 1.05, 0.9, 2.35, 1.25, 0.55])
    else:
        add_callout(doc, "VM discovery", "No virtual machines were returned by the read-only Azure CLI query. Verify subscription scope and permissions before concluding that no VMs exist.", fill=WARNING)

    # Network discovery.
    doc.add_heading("5. Network discovery findings", level=1)
    net_rows: list[list[str]] = []
    for vnet in d.vnets:
        net_rows.append([
            "Virtual Network",
            f"{vnet.get('name', 'Unknown')} / {vnet.get('location', 'Unknown')}",
            ", ".join(q(vnet, "addressSpace.addressPrefixes", []) or []) or "Not verified",
        ])
    if d.vnet_detail:
        subnets = q(d.vnet_detail, "subnets", []) or []
        for s in subnets:
            net_rows.append([
                "Subnet",
                safe_text(s.get("name")),
                safe_text(s.get("addressPrefix")),
            ])
        dns = q(d.vnet_detail, "dhcpOptions.dnsServers", []) or []
        if dns:
            net_rows.append(["VNet DNS", "Configured DNS servers", ", ".join(dns)])
    for key, label in [
        ("routeTables", "Route tables"),
        ("virtualNetworkGateways", "Virtual network gateways"),
        ("azureFirewalls", "Azure Firewall"),
        ("natGateways", "NAT gateways"),
        ("applicationGateways", "Application Gateways"),
        ("expressRouteCircuits", "ExpressRoute circuits"),
        ("connections", "VPN/connection objects"),
        ("localNetworkGateways", "Local network gateways"),
    ]:
        value = d.subscription_resource_checks.get(key)
        if value is None:
            obs = "Not verified"
        elif value:
            names = ", ".join(str(x.get("name")) for x in value[:8])
            obs = names + (f" (+{len(value)-8} more)" if len(value) > 8 else "")
        else:
            obs = "None returned by subscription-wide discovery"
        net_rows.append([label, "Subscription-wide discovery", obs])
    add_table(doc, ["Finding", "Scope", "Observed state"], net_rows, [1.5, 2.1, 4.4])

    doc.add_paragraph("Private DNS zones:")
    dns_rows = [[z.get("name", "Unknown"), z.get("location", "global")] for z in d.private_dns_zones]
    add_table(doc, ["Zone", "Location"], dns_rows or [["None returned", ""]], [5.8, 1.2])

    # Security.
    doc.add_heading("6. Existing security exposure observed", level=1)
    security_rows: list[list[str]] = []
    for nsg_name, rules in d.nsg_rules.items():
        for rule in rules:
            direction = str(rule.get("direction", ""))
            access = str(rule.get("access", ""))
            if direction.lower() == "inbound" and access.lower() == "allow":
                ports = rule.get("destinationPortRange") or rule.get("destinationPortRanges") or "*"
                src = rule.get("sourceAddressPrefix") or rule.get("sourceAddressPrefixes") or rule.get("sourceApplicationSecurityGroups") or "*"
                security_rows.append([
                    str(rule.get("name", "Unknown")),
                    str(src),
                    str(ports),
                    f"{nsg_name}: existing inbound allow",
                ])
    if not security_rows:
        security_rows.append(["No explicit allow rules returned", "", "", "Review effective rules on target NICs before concluding exposure."])
    add_table(doc, ["Rule", "Source", "Port", "Assessment"], security_rows[:12], [1.6, 2.25, 1.0, 2.15])
    add_callout(
        doc,
        "Security discussion point",
        "Existing security controls belong to the current workload environment. Do not copy existing inbound rules into the new East US network without design review. For the new project, prefer private VM addresses and controlled administrative access such as Azure Bastion rather than broad Internet-facing RDP."
    )

    # Proposed direction.
    doc.add_heading("7. Proposed direction for the new project", level=1)
    doc.add_paragraph(
        "This is the working architecture for discussion, not a final build specification. "
        "The client still needs to approve the address plan, Cisco integration details, workload sizing, and security model."
    )
    proposed = [
        ["Azure region", region],
        ["Existing subscription", "Use the existing subscription for the new project while preserving existing West US 3 resources."],
        ["Virtual WAN", "Standard"],
        ["Virtual Hub", f"One {region} hub; hub CIDR pending client/IPAM approval."],
        ["Existing VNet", "Do not modify VN-Accelevation / existing West US 3 network."],
        ["New VNet", "Create a new East US VNet after the client approves a non-overlapping CIDR."],
        ["Cisco", "Two Catalyst 8000V SD-WAN NVAs through the supported Azure Virtual WAN integration path."],
        ["Azure Firewall", "Pending client security/egress decision."],
    ]
    add_table(doc, ["Component", "Working proposal"], proposed, [2.0, 4.9])

    # Architecture image.
    asset_dir = output_path.parent / ".azure_discovery_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    fig_path = asset_dir / "architecture.png"
    architecture_figure(fig_path, current_vnet_space, "East US VNet CIDR pending", "East US Hub CIDR pending")
    if fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fig_path), width=Inches(6.8))
        p = doc.add_paragraph("Figure 1. Working target architecture for the new East US project.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = "Caption"

    doc.add_heading("8. Conceptual target state", level=1)
    add_table(doc, ["Layer", "Working design", "Purpose"], [
        ["Existing West US 3", f"{current_vnet_name} / {current_vnet_space}", "Existing AVD, identity, storage/FSLogix, SQL and support services."],
        ["Connectivity", "Azure Virtual WAN Standard", "Managed transit fabric for the new project."],
        ["Virtual Hub", f"{region}; CIDR pending", "Transit and Cisco NVA integration boundary."],
        ["SD-WAN", "2 x Catalyst 8000V", "Cisco SD-WAN cloud edge and WAN connectivity."],
        ["East US VNet", "New, non-overlapping CIDR", "ISE, identity, PKI, management and workload services."],
    ], [1.5, 2.45, 2.95])

    doc.add_heading("9. Decisions required from the client", level=1)
    add_bullets(doc, [
        f"Confirm {region} as the deployment region.",
        "Provide the corporate, data center, branch, and Cisco SD-WAN IP ranges so the Azure CIDR can be selected without overlap.",
        "Approve the East US Virtual WAN hub CIDR and workload VNet CIDR.",
        "Confirm whether the new East US workloads must communicate with the existing West US 3 environment.",
        "Provide Cisco SD-WAN Manager, Validator, and Catalyst SD-WAN software versions.",
        "Confirm Cisco licensing model: BYOL or Azure Marketplace, plus required throughput tier.",
        "Provide sizing for ISE nodes, identity/DNS, PKI, and all requested VMs.",
        "Confirm Azure Firewall requirement and the intended Internet egress/inspection model.",
        "Confirm backup, logging, monitoring, and retention requirements for the new workloads.",
        "Confirm ongoing Azure, network, and security ownership after deployment.",
    ])

    doc.add_heading("10. Immediate next steps", level=1)
    next_rows = [
        ["1", "Lock the client-approved East US CIDR and Virtual WAN hub CIDR.", "Client"],
        ["2", "Complete Cisco version/licensing/throughput discovery.", "Client"],
        ["3", "Create new East US resource groups for the project.", "Engineering"],
        ["4", "Create Azure Virtual WAN Standard and East US Virtual Hub.", "Engineering"],
        ["5", "Create the new East US VNet and approved subnets.", "Engineering"],
        ["6", "Deploy Cisco Catalyst 8000V through the supported Azure Virtual WAN integration.", "Engineering / Cisco"],
        ["7", "Connect workloads, implement security/monitoring, and validate routing.", "Engineering"],
    ]
    add_table(doc, ["Step", "Action", "Owner"], next_rows, [0.55, 5.35, 1.0])

    doc.add_heading("Appendix A - Current Azure facts from discovery", level=1)
    facts = [
        ["Tenant", safe_text(tenant_name)],
        ["Tenant ID", safe_text(tenant_id)],
        ["Subscription", safe_text(q(d.account, "name"))],
        ["Subscription ID", safe_text(sub_id)],
        ["Existing VNet", safe_text(current_vnet_name)],
        ["Existing VNet region", safe_text(current_vnet_location)],
        ["Existing VNet CIDR", safe_text(current_vnet_space)],
        ["Existing subnet(s)", ", ".join(str(x.get("name")) for x in (q(d.vnet_detail, "subnets", []) or [])) or "Not verified"],
        ["Existing DNS servers", ", ".join(q(d.vnet_detail, "dhcpOptions.dnsServers", []) or []) or "Not verified"],
        ["Existing NSGs", ", ".join(str(x.get("name")) for x in d.nsgs) or "Not verified"],
        ["Existing public IPs", ", ".join(f"{x.get('name')}={x.get('ipAddress')}" for x in d.public_ips) or "None returned"],
        ["Existing private endpoints", ", ".join(str(x.get("name")) for x in d.private_endpoints) or "None returned"],
        ["Existing private DNS zones", ", ".join(str(x.get("name")) for x in d.private_dns_zones) or "None returned"],
    ]
    add_table(doc, ["Item", "Current value"], facts, [2.1, 4.8])

    doc.add_paragraph("Virtual machine inventory summary:")
    vm_fact_rows = []
    for vm in sorted(d.vms, key=lambda x: str(x.get("name", "")).lower()):
        name = str(vm.get("name", "Unknown"))
        detail = d.vm_details.get(name, {})
        vm_fact_rows.append([
            name,
            safe_text(vm.get("resourceGroup")),
            safe_text(vm.get("location")),
            safe_text(vm.get("powerState")),
            safe_text(q(detail, "hardwareProfile.vmSize", vm.get("hardwareProfile", {}).get("vmSize"))),
        ])
    add_table(doc, ["VM", "Resource group", "Region", "State", "Size"], vm_fact_rows or [["None returned", "", "", "", ""]], [1.55, 1.8, 1.0, 1.15, 1.4])

    if d.errors:
        doc.add_paragraph("Discovery exceptions:", style="Heading 2")
        add_bullets(doc, d.errors[:20])

    doc.add_paragraph(
        "Source note: Values in this report are generated from read-only Azure CLI discovery at run time. "
        "Items not returned by a successful subscription-wide query are reported as none returned; failed queries are reported as discovery exceptions."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a client-facing Azure current-state discovery DOCX.")
    parser.add_argument("--subscription", help="Azure subscription ID. If omitted, uses the current az context.")
    parser.add_argument("--output", default="Azure_Current_State_and_Project_Discovery.docx")
    parser.add_argument("--region", default="East US", help="Proposed new-project region.")
    args = parser.parse_args()

    if shutil.which("az") is None:
        print("ERROR: Azure CLI (az) is not installed or not in PATH.", file=sys.stderr)
        return 2

    d = discover(args.subscription)
    output = Path(args.output).resolve()
    build_report(d, output, args.region)

    print(f"Generated: {output}")
    if d.errors:
        print(f"Discovery completed with {len(d.errors)} exception(s). See report Appendix A.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
